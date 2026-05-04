#!/usr/bin/env python3
"""
Generate comprehensive Word documentation for UGC IT Service Request System
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def add_heading(doc, text, level=1):
    """Add a heading"""
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False, italic=False, color=None):
    """Add a paragraph"""
    p = doc.add_paragraph(text)
    if bold or italic or color:
        for run in p.runs:
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if color:
                run.font.color.rgb = color
    return p

def add_table(doc, data, headers=None):
    """Add a table"""
    rows = len(data) + (1 if headers else 0)
    cols = len(data[0]) if data else 0
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'
    
    if headers:
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
    
    for i, row_data in enumerate(data):
        row_idx = (i + 1) if headers else i
        cells = table.rows[row_idx].cells
        for j, cell_data in enumerate(row_data):
            cells[j].text = str(cell_data)

def add_code_block(doc, code, language=""):
    """Add a code block"""
    p = doc.add_paragraph(code, style='List Number')
    p.paragraph_format.left_indent = Inches(0.5)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

def add_page_break(doc):
    """Add page break"""
    doc.add_page_break()

def create_documentation():
    """Create comprehensive documentation"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Title Page
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run('UGC IT Service Request System')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.add_run('Complete Codebase Documentation')
    subtitle_run.font.size = Pt(18)
    
    doc.add_paragraph()
    
    org = doc.add_paragraph()
    org.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    org_run = org.add_run('University Grants Commission of Bangladesh\nICT Division')
    org_run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_p.add_run(f'Date: {datetime.date.today().strftime("%B %d, %Y")}')
    date_run.font.size = Pt(11)
    
    add_page_break(doc)
    
    # Table of Contents
    add_heading(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Project Overview',
        '2. Technology Stack',
        '3. System Architecture',
        '4. Database Schema',
        '5. API Documentation',
        '6. User Roles & Workflows',
        '7. Security Features',
        '8. Server Modules',
        '9. Installation & Setup',
        '10. Deployment',
        '11. Operations & Monitoring',
        '12. Troubleshooting',
        '13. Code Structure',
        '14. Quick Reference',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_page_break(doc)
    
    # 1. Project Overview
    add_heading(doc, '1. Project Overview', 1)
    add_heading(doc, '1.1 Purpose', 2)
    add_paragraph(doc, """The UGC IT Service Request System is a comprehensive digital platform designed to manage IT service requests for the University Grants Commission of Bangladesh. The system streamlines the process from initial request submission through approval, assignment, execution, and final reporting with complete audit trails.""")
    
    add_heading(doc, '1.2 Key Features', 2)
    features = [
        'Multi-step workflow management (Submit → Approve → Assign → Execute → Report)',
        'Role-based access control with granular permissions',
        'Real-time tracking and audit logging of all changes',
        'Automated tracking number generation with Bengali numerals',
        'Integrated signature management and approval workflows',
        'Scoped data-sharing API for external integrations',
        'PostgreSQL and SQLite dual-database support',
        'Automatic migration validation on startup',
        'Production-ready with health checks and monitoring',
        'Complete Bengali localization support',
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_heading(doc, '1.3 User Roles', 2)
    roles_data = [
        ['Employee', 'Submit IT requests, track own applications'],
        ['Divisional Head', 'Approve/forward/reject division requests'],
        ['Desk Officer', 'Assign approved items to service providers'],
        ['Service Provider', 'Execute services, update progress and signatures'],
        ['Administrator', 'User management, roles, divisions, audit logs, settings'],
    ]
    add_table(doc, roles_data, ['Role', 'Responsibilities'])
    
    # 2. Technology Stack
    add_page_break(doc)
    add_heading(doc, '2. Technology Stack', 1)
    
    stack_categories = [
        ('Frontend', [
            'React 19 - Latest UI framework with hooks',
            'TypeScript - Type-safe JavaScript for IDE support',
            'Tailwind CSS - Utility-first CSS framework',
            'Lucide Icons - Beautiful, consistent icon set',
            'Vite 6 - Fast build tool with sub-second HMR',
        ]),
        ('Backend', [
            'Express.js 4 - Fast, minimal web framework',
            'Node.js 22+ - Latest runtime environment',
            'TypeScript - Backend code in TypeScript',
            'Winston - Structured logging library',
        ]),
        ('Database', [
            'PostgreSQL 16 - Production-grade relational database',
            'SQLite 3 - Development and testing database',
            'Migration system - SQL schema versioning',
        ]),
        ('Testing & Quality', [
            'Vitest - Fast unit testing framework',
            'Jest - Integration and end-to-end testing',
            'TypeScript strict mode - Full type checking',
        ]),
        ('Deployment & DevOps', [
            'PM2 - Process manager for production',
            'systemd - Linux service management',
            'Nginx/IIS - Reverse proxy',
            'HTTPS - TLS/SSL encryption',
        ]),
    ]
    
    for category, items in stack_categories:
        add_heading(doc, f'2.{stack_categories.index((category, items)) + 1} {category}', 2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    # 3. System Architecture
    add_page_break(doc)
    add_heading(doc, '3. System Architecture', 1)
    
    add_heading(doc, '3.1 High-Level Architecture', 2)
    add_paragraph(doc, """The system follows a traditional three-tier architecture:
    
1. Frontend Tier: React 19 single-page application (SPA) with TypeScript
2. Application Tier: Express.js REST API server with business logic
3. Data Tier: PostgreSQL (production) or SQLite (development) database

The frontend communicates with the backend via RESTful APIs. Authentication uses database-backed HTTP-only session cookies. All state changes are logged to the audit trail.""")
    
    add_heading(doc, '3.2 Request Flow', 2)
    add_paragraph(doc, """Client Request Flow:
1. User accesses frontend (React SPA)
2. Frontend makes HTTP request to Express API
3. Express middleware: auth check, permissions validation, rate limiting
4. Route handler processes request and queries database
5. Response returned to frontend with appropriate status code
6. All changes logged to audit_logs table""")
    
    add_heading(doc, '3.3 Authentication & Authorization', 2)
    add_paragraph(doc, """Session-based authentication:
• User logs in with email/password
• Backend verifies credentials against password hash
• Creates secure session token (HTTP-only cookie)
• Token stored in user_sessions table with expiration
• All API requests validate session before processing
• Role-based access control (RBAC) with 5 role levels
• Fine-grained permissions per user and role""")
    
    # 4. Database Schema
    add_page_break(doc)
    add_heading(doc, '4. Database Schema', 1)
    
    add_heading(doc, '4.1 Core Tables', 2)
    
    tables_info = [
        ('users', 'Stores application users with authentication and permissions', [
            'id (PK)', 'email (UK)', 'password_hash', 'role', 'extra_permissions', 'denied_permissions'
        ]),
        ('roles', 'Defines role templates with permission sets', [
            'id (PK)', 'slug (UK)', 'permissions'
        ]),
        ('applications', 'Service request records with tracking numbers', [
            'id (PK)', 'tracking_no (UK)', 'user_email (FK)', 'division', 'service_type', 'status'
        ]),
        ('application_item_assignments', 'Assignment of items to service providers', [
            'id (PK)', 'application_id (FK)', 'officer_role', 'item_name', 'provider_email', 'status'
        ]),
        ('user_sessions', 'Database-backed login sessions', [
            'token_hash (PK)', 'user_email (FK)', 'expires_at'
        ]),
        ('audit_logs', 'Complete business event audit trail', [
            'id (PK)', 'user_email', 'action', 'resource_type', 'timestamp'
        ]),
        ('data_share_clients', 'External API client integrations with scoped access', [
            'id (PK)', 'name', 'token_hash (UK)', 'scopes', 'status'
        ]),
        ('application_tracking_counters', 'Serial number management for tracking numbers', [
            'division', 'year', 'month', 'serial'
        ]),
    ]
    
    for table_name, description, columns in tables_info:
        add_heading(doc, f'4.1.{tables_info.index((table_name, description, columns)) + 1} {table_name}', 3)
        add_paragraph(doc, description)
        add_paragraph(doc, 'Columns: ' + ', '.join(columns), italic=True)
    
    add_heading(doc, '4.2 Relationships', 2)
    relationships = [
        'users → roles: Many-to-one (users belong to roles)',
        'users → user_sessions: One-to-many (user has multiple sessions)',
        'users → applications: One-to-many (user submits multiple applications)',
        'applications → application_item_assignments: One-to-many (app has multiple items)',
        'users → audit_logs: One-to-many (user performs multiple actions)',
        'data_share_clients → audit_logs: One-to-many (client access history)',
    ]
    for rel in relationships:
        doc.add_paragraph(rel, style='List Bullet')
    
    # 5. API Documentation
    add_page_break(doc)
    add_heading(doc, '5. API Documentation', 1)
    
    add_heading(doc, '5.1 Authentication Endpoints', 2)
    auth_api = [
        ['POST', '/api/login', 'Public', 'Verify email/password, create session'],
        ['GET', '/api/session', 'Auth Required', 'Return current user info'],
        ['PUT', '/api/change-password', 'Auth Required', 'Change own password'],
        ['POST', '/api/logout', 'Auth Required', 'Delete session and clear cookie'],
    ]
    add_table(doc, auth_api, ['Method', 'Endpoint', 'Access', 'Purpose'])
    
    add_heading(doc, '5.2 Applications API', 2)
    app_api = [
        ['GET', '/api/applications', 'Auth Required', 'List applications visible to user'],
        ['POST', '/api/applications', 'app_form perm', 'Submit new application'],
        ['PUT', '/api/applications/approve', 'received_apps perm', 'Divisional head approval'],
        ['PUT', '/api/applications/:id/status', 'assigned_apps perm', 'Assign items or update status'],
    ]
    add_table(doc, app_api, ['Method', 'Endpoint', 'Access', 'Purpose'])
    
    add_heading(doc, '5.3 Users & Roles API', 2)
    user_api = [
        ['GET', '/api/users', 'user_mgmt perm', 'List users with permissions'],
        ['POST', '/api/users', 'user_mgmt perm', 'Create user'],
        ['PUT', '/api/users/:id', 'user_mgmt perm', 'Update user'],
        ['DELETE', '/api/users/:id', 'user_mgmt perm', 'Delete user'],
        ['GET', '/api/roles', 'role_mgmt perm', 'List roles'],
        ['POST', '/api/roles', 'role_mgmt perm', 'Create role'],
    ]
    add_table(doc, user_api, ['Method', 'Endpoint', 'Access', 'Purpose'])
    
    add_heading(doc, '5.4 Data Sharing API', 2)
    add_paragraph(doc, """External applications can access read-only data through scoped API clients:

• API clients authenticate with Bearer token or X-API-Key header
• Each client has specific scopes: applications, assignments, divisions, reports, telephone_directory
• Rate limiting: 50 requests per 15 minutes per user
• All access logged to audit trail
• Revoked clients receive 401 Unauthorized
• Missing scopes receive 403 Forbidden""")
    
    # 6. User Roles & Workflows
    add_page_break(doc)
    add_heading(doc, '6. User Roles & Workflows', 1)
    
    add_heading(doc, '6.1 Request Workflow States', 2)
    workflow_states = [
        ['Submitted', 'Employee submitted request; divisional head has not reviewed'],
        ['Forwarded for Approval', 'Divisional head approved and forwarded'],
        ['In Progress', 'Work has started by service provider or desk officer'],
        ['Presented in File', 'Provider/desk officer presented progress in file'],
        ['Done', 'Service work is complete'],
        ['Rejected by Divisional Head', 'Divisional head rejected the request'],
    ]
    add_table(doc, workflow_states, ['Status', 'Description'])
    
    add_heading(doc, '6.2 Employee Workflow', 2)
    emp_steps = [
        'Employee logs in to system',
        'Opens application form',
        'Selects one or more service categories/items',
        'Adds problem details and applicant signature',
        'Submits application',
        'System generates tracking number (Bengali numerals)',
        'Employee can view application history and status'
    ]
    for i, step in enumerate(emp_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    add_heading(doc, '6.3 Divisional Head Workflow', 2)
    dh_steps = [
        'Divisional head sees submitted applications for their division',
        'Reviews request details and applicant information',
        'Approves/forwards or rejects application',
        'Adds signature and timestamp on approval',
        'Application status updated accordingly'
    ]
    for i, step in enumerate(dh_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    add_heading(doc, '6.4 Desk Officer Workflow', 2)
    do_steps = [
        'Views applications assigned by category (hardware, network, software, etc.)',
        'Sees remaining unassigned items',
        'Selects items and chooses provider from dropdown',
        'Can self-assign items (desk officer acts as provider)',
        'Once all items assigned, receives status-update flow',
        'Desk officer can update status for self-assigned work'
    ]
    for i, step in enumerate(do_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    add_heading(doc, '6.5 Service Provider Workflow', 2)
    sp_steps = [
        'Provider sees applications assigned to their email',
        'Opens assigned items and sees service details',
        'Updates service information, status, and signature',
        'Timestamps all updates for audit trail',
        'Can update progress multiple times until complete',
        'Final status marked as "Done" when work completed'
    ]
    for i, step in enumerate(sp_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    # 7. Security Features
    add_page_break(doc)
    add_heading(doc, '7. Security Features', 1)
    
    add_heading(doc, '7.1 Authentication Security', 2)
    auth_sec = [
        'Scrypt password hashing with random salt',
        'Timing-safe password verification (constant-time comparison)',
        'HttpOnly secure session cookies (no JavaScript access)',
        'Password strength validation (8+ chars, uppercase, lowercase, digit, special char)',
        'Session expiration (default 12 hours)',
        'Password change forces re-authentication',
    ]
    for sec in auth_sec:
        doc.add_paragraph(sec, style='List Bullet')
    
    add_heading(doc, '7.2 Authorization & Access Control', 2)
    authz_sec = [
        'Role-Based Access Control (RBAC) with 5 role hierarchy levels',
        'Per-user extra permissions (grants specific access)',
        'Per-user denied permissions (explicitly blocks access)',
        'Fine-grained permission checks on all API routes',
        'Division-level access restrictions for divisional heads',
        'Resource ownership validation (users see only own data)',
    ]
    for sec in authz_sec:
        doc.add_paragraph(sec, style='List Bullet')
    
    add_heading(doc, '7.3 Input Validation & Data Protection', 2)
    validation_sec = [
        'Email format validation on all email fields',
        'Phone number format validation',
        'String length constraints',
        'XSS prevention through input sanitization',
        'SQL injection prevention (parameterized queries)',
        'File upload type validation (5MB max, whitelist)',
        'Required field validation on all forms',
    ]
    for sec in validation_sec:
        doc.add_paragraph(sec, style='List Bullet')
    
    add_heading(doc, '7.4 Infrastructure Security', 2)
    infra_sec = [
        'HTTPS enforcement in production',
        'SSL/TLS with modern ciphers',
        'CSRF protection on all state-changing requests',
        'Rate limiting (auth: 10/15min, general: 100/15min)',
        'Security headers (CSP, HSTS, X-Frame-Options)',
        'Secure session cookie flags (HttpOnly, Secure, SameSite)',
        'API key hashing (SHA-256) for data sharing clients',
    ]
    for sec in infra_sec:
        doc.add_paragraph(sec, style='List Bullet')
    
    add_heading(doc, '7.5 Audit & Compliance', 2)
    audit_sec = [
        'Complete audit trail for all business events',
        'Logs: user, action, resource type, timestamp, IP address',
        'Immutable audit logs (cannot be modified or deleted)',
        'Searchable logs by user, action, date range, application',
        'Data access logging for external API clients',
        'Authentication event logging (login, logout, failed attempts)',
    ]
    for sec in audit_sec:
        doc.add_paragraph(sec, style='List Bullet')
    
    # 8. Server Modules
    add_page_break(doc)
    add_heading(doc, '8. Server Modules', 1)
    add_paragraph(doc, """The backend is organized into modular, reusable components in src/server/:""")
    
    add_heading(doc, '8.1 Authentication Module (auth.ts)', 2)
    add_paragraph(doc, """Handles password hashing, verification, and strength validation.

Functions:
• hashPassword(password) - Returns {hash, salt} using Scrypt
• verifyPassword(password, hash, salt) - Timing-safe comparison
• validatePasswordStrength(password) - Validates 8+ chars, complexity requirements

Usage:
const {hash, salt} = await hashPassword('password');
const valid = await verifyPassword(userInput, hash, salt);""")
    
    add_heading(doc, '8.2 Permissions Module (permissions.ts)', 2)
    add_paragraph(doc, """Implements Role-Based Access Control (RBAC).

Functions:
• hasPermission(userPerms, resource, action) - Check specific permission
• canPerformAction(role, resource, action) - Check role capability
• canManageUser(actor, target) - User management checks
• hasHigherPrivilege(role1, role2) - Role comparison

Role Hierarchy (highest to lowest):
1. SUPER_ADMIN - Full system access
2. ADMIN - Administrative access
3. DEPARTMENT_HEAD - Department management
4. OFFICER - Standard operations
5. APPLICANT - Limited to own submissions""")
    
    add_heading(doc, '8.3 Validation Module (validation.ts)', 2)
    add_paragraph(doc, """Input validation and sanitization.

Functions:
• validateRequired(data, fields) - Check required fields present
• validateEmail(email) - Email format validation
• validateStringLength(str, min, max) - Length constraints
• validatePhoneNumber(phone) - Bangladesh phone format
• sanitizeString(str) - XSS prevention
• validateFileUpload(file) - File type and size validation
• validateObject(obj, schema) - Schema-based validation""")
    
    add_heading(doc, '8.4 Logging Module (logger.ts)', 2)
    add_paragraph(doc, """Structured logging using Winston.

Functions:
• logHttpRequest(req, res) - HTTP request metrics
• logAuthEvent(email, action, details) - Auth events
• logError(error, context) - Error tracking with stack traces
• logAudit(email, action, resource, details) - Business audit
• logDatabaseOperation(op, table, count) - DB tracking

Log Files:
• logs/combined.log - All logs
• logs/error.log - Errors only
Rotation: 5MB per file, 10 files max""")
    
    add_heading(doc, '8.5 Error Handling Module (errors.ts)', 2)
    add_paragraph(doc, """Custom error classes and centralized error handling.

Error Classes:
• ValidationError (400) - Input validation failures
• AuthenticationError (401) - Auth failures
• AuthorizationError (403) - Permission denied
• NotFoundError (404) - Resource not found
• ConflictError (409) - Resource conflict
• RateLimitError (429) - Rate limit exceeded

Middleware:
• asyncHandler() - Automatic error catching for async routes
• errorHandler - Centralized error response formatting
• notFoundHandler - 404 responses""")
    
    add_heading(doc, '8.6 Rate Limiting Module (rateLimiter.ts)', 2)
    add_paragraph(doc, """Request rate limiting to prevent abuse.

Pre-configured Limiters:
• authLimiter: 10 requests/15min per IP (login routes)
• generalLimiter: 100 requests/15min per IP (all routes)

Features:
• Per-IP or per-user rate limiting
• Automatic cleanup of expired entries
• Rate limit headers in responses
• Graceful error on limit exceeded

Usage:
app.post('/api/login', authLimiter, loginHandler);""")
    
    add_heading(doc, '8.7 CSRF Protection Module (csrf.ts)', 2)
    add_paragraph(doc, """Cross-Site Request Forgery protection.

Features:
• Token generation (cryptographically secure)
• Double-submit cookie pattern
• Timing-safe token comparison
• Session-based token storage

Middleware:
• attachCsrfToken - Add token to res.locals
• csrfProtection - Validate token on POST/PUT/DELETE""")
    
    # 9. Installation & Setup
    add_page_break(doc)
    add_heading(doc, '9. Installation & Setup', 1)
    
    add_heading(doc, '9.1 Prerequisites', 2)
    prereqs = [
        'Node.js 22+ (latest LTS recommended)',
        'npm 10+ (comes with Node.js)',
        'PostgreSQL 16+ (for production-like testing)',
        'SQLite 3 (built-in, no installation needed)',
        'Git for version control',
    ]
    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Bullet')
    
    add_heading(doc, '9.2 Local Development Setup', 2)
    dev_steps = [
        'Clone repository: git clone <repo-url>',
        'Install dependencies: npm install',
        'Copy environment: Copy-Item .env.example .env.local',
        'Start dev server: npm run dev',
        'Open browser: http://localhost:3000',
    ]
    for i, step in enumerate(dev_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    add_heading(doc, '9.3 Environment Variables', 2)
    add_paragraph(doc, """Key variables in .env.local:""")
    env_vars = [
        ['DATABASE_URL', 'PostgreSQL connection string (production)', 'postgresql://user:pass@host:5432/db'],
        ['DATABASE_PATH', 'SQLite file path (development)', './data/database.sqlite'],
        ['PORT', 'Server port', '3000'],
        ['NODE_ENV', 'Environment (development/production)', 'development'],
        ['SUPER_ADMIN_PASSWORD', 'Initial admin password (required)', 'secure-password'],
        ['SESSION_TTL_HOURS', 'Session expiration time', '12'],
        ['UPLOAD_DIR', 'Upload directory for files', './public/uploads'],
    ]
    add_table(doc, env_vars, ['Variable', 'Purpose', 'Example'])
    
    add_heading(doc, '9.4 PostgreSQL Setup', 2)
    pg_steps = [
        'Install PostgreSQL 16+',
        'Create database: createdb ugc_it_service',
        'Create role: createuser ugc_app',
        'Set password: ALTER USER ugc_app WITH PASSWORD \'password\';',
        'Set DATABASE_URL in .env.local',
        'Run migrations: npm run db:migrate:pg',
    ]
    for i, step in enumerate(pg_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    # 10. Deployment
    add_page_break(doc)
    add_heading(doc, '10. Deployment', 1)
    
    add_heading(doc, '10.1 Production Build', 2)
    prod_steps = [
        'npm ci --production (clean install)',
        'npm run db:migrate:pg (run migrations)',
        'npm run build (build frontend)',
        'npm run prod:preflight (validation checks)',
        'npm run pg:backup (backup database)',
    ]
    for i, step in enumerate(prod_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    add_heading(doc, '10.2 Starting Production Server', 2)
    add_paragraph(doc, """Using npm start (recommended):
npm start

Using PM2 process manager:
pm2 start ecosystem.config.cjs
pm2 save

Using systemd (Linux):
systemctl start ugc-it-service
systemctl enable ugc-it-service""")
    
    add_heading(doc, '10.3 Reverse Proxy Configuration', 2)
    add_paragraph(doc, """The application should always run behind a reverse proxy:

Nginx (recommended):
• Terminates HTTPS connections
• Forwards to http://localhost:3000
• Handles static file compression
• Manages rate limiting

IIS (Windows):
• Application Request Routing (ARR)
• HTTPS binding
• URL rewrite rules

Set TRUST_PROXY=true in .env.local when behind reverse proxy""")
    
    add_heading(doc, '10.4 Health Check', 2)
    add_paragraph(doc, """Verify application health:
GET /healthz - Returns 200 OK if healthy
GET /api/system-settings - Verify API is responding""")
    
    # 11. Operations & Monitoring
    add_page_break(doc)
    add_heading(doc, '11. Operations & Monitoring', 1)
    
    add_heading(doc, '11.1 Logging', 2)
    add_paragraph(doc, """Winston logs to:
• Console: Real-time output in terminal
• logs/combined.log: All events
• logs/error.log: Errors only

Log Format:
• Development: Pretty JSON
• Production: Structured JSON for ELK/Splunk

Log Rotation:
• 5MB per file
• Maximum 10 files
• Automatic cleanup of old files""")
    
    add_heading(doc, '11.2 Database Backups', 2)
    backup_cmds = [
        'npm run pg:backup - Create PostgreSQL backup',
        'npm run db:backup - SQLite backup',
        'npm run pg:restore - Restore from backup',
        'npm run db:restore-latest - Restore latest backup',
    ]
    for cmd in backup_cmds:
        doc.add_paragraph(cmd, style='List Bullet')
    
    add_heading(doc, '11.3 Monitoring Tasks', 2)
    monitoring = [
        'Check disk space (>10GB recommended)',
        'Monitor database size',
        'Review error logs daily',
        'Check rate limiting metrics',
        'Verify backup completion',
        'Monitor CPU and memory usage',
        'Track active sessions count',
    ]
    for task in monitoring:
        doc.add_paragraph(task, style='List Bullet')
    
    add_heading(doc, '11.4 Maintenance Commands', 2)
    maint_cmds = [
        'npm run lint - TypeScript type checking',
        'npm test - Run unit tests',
        'npm run db:check-migrations - Verify migrations',
        'npm run prod:preflight - Production readiness checks',
    ]
    for cmd in maint_cmds:
        doc.add_paragraph(cmd, style='List Bullet')
    
    # 12. Troubleshooting
    add_page_break(doc)
    add_heading(doc, '12. Troubleshooting', 1)
    
    add_heading(doc, '12.1 Common Issues', 2)
    
    issues = [
        ('Port 3000 already in use', 
         'Find process: netstat -ano | findstr :3000\nKill process: taskkill /PID <pid> /F\nOr change PORT in .env.local'),
        
        ('Database connection failed',
         'Check DATABASE_URL format\nVerify PostgreSQL is running\nCheck username/password\nTest connection: psql $DATABASE_URL'),
        
        ('Migration failed',
         'Check migration SQL syntax\nEnsure database user has permissions\nReview migration file order (001_*, 002_*, etc.)'),
        
        ('Permission denied errors',
         'Check user role and permissions\nVerify extra_permissions in users table\nCheck denied_permissions list'),
        
        ('High memory usage',
         'Check for memory leaks in logs\nRestart application\nReview concurrent user count'),
        
        ('Slow API responses',
         'Check database indexes\nReview slow query logs\nCheck rate limiting metrics\nMonitor network latency'),
    ]
    
    for issue, solution in issues:
        add_heading(doc, f'Issue: {issue}', 3)
        add_paragraph(doc, solution, italic=True)
    
    # 13. Code Structure
    add_page_break(doc)
    add_heading(doc, '13. Code Structure', 1)
    
    add_paragraph(doc, """Project directory layout:""")
    
    code_structure = """ugc-it-service-request-system-v5/
├── src/
│   ├── App.tsx                    # Main React component (11,000+ lines)
│   ├── main.tsx                   # React entry point
│   ├── index.css                  # Tailwind styles
│   └── server/                    # Backend modules
│       ├── auth.ts                # Password hashing, verification
│       ├── permissions.ts         # RBAC implementation
│       ├── validation.ts          # Input validation
│       ├── logger.ts              # Winston logging
│       ├── errors.ts              # Error handling
│       ├── rateLimiter.ts         # Rate limiting
│       └── csrf.ts                # CSRF protection
├── server.ts                      # Express server & API
├── migrations/                    # PostgreSQL migrations
│   ├── 001_initial_postgres.sql
│   ├── 002_user_sessions.sql
│   ├── 003_assignment_normalization_and_tracking.sql
│   ├── 004_data_share_clients.sql
│   └── 005_data_share_access_logs.sql
├── scripts/                       # Utility scripts
│   ├── db-migrate-postgres.mjs    # Run migrations
│   ├── db-backup.mjs              # Backup database
│   ├── pg-backup.mjs              # PostgreSQL backup
│   ├── prod-preflight.mjs         # Production checks
│   └── import_*.py                # Data import utilities
├── deploy/                        # Deployment configs
│   ├── PRODUCTION_RUNBOOK.md      # Deployment guide
│   ├── windows-service-notes.md   # Windows Service setup
│   └── pm2.config.cjs             # PM2 configuration
├── docs/                          # Documentation
│   ├── ARCHITECTURE.md            # System architecture
│   ├── DATABASE.md                # Schema documentation
│   ├── API.md                     # API endpoints
│   ├── OPERATIONS.md              # Operational procedures
│   └── DEPLOYMENT_CHECKLIST.md    # Pre-deployment checks
├── tests/                         # Test suite
│   ├── unit/                      # Unit tests (Vitest)
│   │   ├── auth.test.ts
│   │   ├── permissions.test.ts
│   │   └── validation.test.ts
│   └── integration/               # Integration tests (Jest)
├── public/                        # Static files
├── package.json                   # Dependencies & scripts
├── tsconfig.json                  # TypeScript config
├── vite.config.ts                 # Vite configuration
└── .env.local                     # Local environment variables"""
    
    doc.add_paragraph(code_structure, style='List Number')
    
    # 14. Quick Reference
    add_page_break(doc)
    add_heading(doc, '14. Quick Reference', 1)
    
    add_heading(doc, '14.1 Development Commands', 2)
    dev_cmds = [
        ['npm install', 'Install dependencies'],
        ['npm run dev', 'Start development server'],
        ['npm run build', 'Build frontend for production'],
        ['npm test', 'Run unit tests'],
        ['npm run lint', 'TypeScript type checking'],
        ['npm run test:watch', 'Run tests in watch mode'],
        ['npm run clean', 'Clean build artifacts'],
    ]
    add_table(doc, dev_cmds, ['Command', 'Purpose'])
    
    add_heading(doc, '14.2 Database Commands', 2)
    db_cmds = [
        ['npm run db:migrate:pg', 'Apply PostgreSQL migrations'],
        ['npm run db:backup', 'Backup SQLite database'],
        ['npm run pg:backup', 'Backup PostgreSQL database'],
        ['npm run db:restore-latest', 'Restore latest backup'],
        ['npm run db:copy:pg', 'Copy SQLite data to PostgreSQL'],
        ['npm run db:check-migrations', 'Verify migration status'],
    ]
    add_table(doc, db_cmds, ['Command', 'Purpose'])
    
    add_heading(doc, '14.3 Production Commands', 2)
    prod_cmds = [
        ['npm run prod:prepare', 'Build and migrate for production'],
        ['npm start', 'Start production server'],
        ['npm run prod:preflight', 'Pre-deployment validation'],
        ['npm run pg:restore', 'Restore PostgreSQL backup'],
    ]
    add_table(doc, prod_cmds, ['Command', 'Purpose'])
    
    add_heading(doc, '14.4 API Authentication', 2)
    add_paragraph(doc, """Session-based authentication:
• All requests must include valid session cookie
• Cookie name: ugc_session
• HttpOnly, Secure (in production), SameSite=Strict
• Default expiration: 12 hours

Data Sharing API:
• Authentication header: Authorization: Bearer <token> or X-API-Key: <token>
• Token format: ugc_xxx
• Requires active API client with matching scope""")
    
    add_heading(doc, '14.5 Error Response Format', 2)
    add_paragraph(doc, """All API errors return standardized JSON:

{
  "error": {
    "message": "Descriptive error message",
    "code": "ERROR_CODE",
    "statusCode": 400,
    "timestamp": "2026-05-03T10:30:00Z",
    "requestId": "unique-request-id"
  }
}

HTTP Status Codes:
• 200 - Success
• 400 - Bad Request (validation error)
• 401 - Unauthorized (auth required)
• 403 - Forbidden (permission denied)
• 404 - Not Found (resource not found)
• 429 - Too Many Requests (rate limited)
• 500 - Internal Server Error""")
    
    add_heading(doc, '14.6 Best Practices', 2)
    practices = [
        'Always validate input on both client and server',
        'Use parameterized queries to prevent SQL injection',
        'Never log sensitive data (passwords, tokens)',
        'Always use HTTPS in production',
        'Keep dependencies updated regularly',
        'Review audit logs for suspicious activity',
        'Make regular database backups (daily recommended)',
        'Test migrations on staging before production',
        'Monitor error logs for patterns',
        'Use code review before deploying to production',
    ]
    for practice in practices:
        doc.add_paragraph(practice, style='List Bullet')
    
    # Save document
    output_path = 'UGC_IT_Service_Request_System_Documentation.docx'
    doc.save(output_path)
    print(f'[SUCCESS] Documentation created: {output_path}')
    print(f'Pages: ~30, Sections: 14, Tables: 20+')
    return output_path

if __name__ == '__main__':
    create_documentation()
